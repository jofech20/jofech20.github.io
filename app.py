# Librerías necesarias para lectura de PDFs, manejo web, IA, y manipulación de texto/documentos
import os
import re
import requests
import pandas as pd
from openai import OpenAI
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import uuid
import math
from collections import Counter

# Inicializa la app Flask
app = Flask(__name__)
CORS(app, origins="https://jofech20.github.io")  # Habilita CORS para la interfaz frontend

# Configuración para subir archivos
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Definición de archivos válidos
ALLOWED_EXTENSIONS = {'pdf'}

# Cliente OpenAI usando la API Key desde variables de entorno
client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))

# API Key de Elsevier para obtener metadatos
API_KEY = os.environ.get('ELSEVIER_API_KEY')

# Carga el dataset de SCImago (revistas y rankings)
SCIMAGO_CSV = 'scimago.csv'
scimago_df = pd.read_csv(SCIMAGO_CSV, sep=';', on_bad_lines='skip')

# Valida que el archivo sea PDF
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Extrae texto de un PDF utilizando PyPDF2
def extract_text_from_pdf(pdf_path):
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ''
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
        print(f"Texto extraído del PDF: {text[:500]}...")
        return text
    except Exception as e:
        print(f"Error al extraer texto: {e}")
        return None

# Extrae el DOI usando una expresión regular desde el texto
def extract_doi_from_text(text):
    match = re.search(r'\b10\.\d{4,9}/[-._;()/:A-Z0-9]+', text, re.IGNORECASE)
    if match:
        doi = match.group(0).strip().replace(' ', '').split('RESEARCH')[0]  # Limpia texto si el DOI es mal capturado
        print(f"DOI encontrado: {doi}")
        return doi
    return None

# Prompt personalizado para generar el estado del arte según estándares científicos
def generate_estado_del_arte(text):
    prompt = f"""
Redacta un **estado del arte** en estilo académico y científico, siguiendo estas indicaciones:

1. Usa un lenguaje claro, objetivo y formal.
2. Estructura el texto con los siguientes subtítulos en negrita usando Markdown:

**Antecedentes del problema**  
(Describe el conocimiento previo sobre el tema, incluyendo enfoques, modelos, teorías o resultados clave que sustentan la investigación.)

**Brechas y vacíos identificados**  
(Indica los aspectos que no han sido resueltos por la literatura previa, limitaciones metodológicas o contradicciones teóricas existentes.)

**Proyección y aporte del artículo analizado**  
(Explica cómo este trabajo contribuye a cerrar brechas previas, qué propone o innova respecto a investigaciones anteriores, y cuál es su valor dentro del campo académico.)

3. Evita repetir el resumen del artículo. Analiza el contexto y redacta desde una perspectiva crítica y sintética.
4. Si es posible, alude a investigaciones mencionadas en el artículo usando frases como “estudios recientes muestran...”, “otros autores han señalado...”, “según la literatura...”. Pero todo según a las citas o referencias que hace el artículo.

Texto base del artículo:
{text[:5000]}
    """
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error al generar estado del arte: {str(e)}"

# Crea el archivo Word incluyendo metadatos y entropía
def save_to_word(estado_arte_texto, filename, metadatos, entropia):
    doc = Document()

    # Encabezado del documento
    doc.add_heading("Estado del Arte", level=1)

    # Metadatos y análisis
    doc.add_paragraph(f"Título del artículo: {metadatos.get('title', 'No disponible')}")
    doc.add_paragraph(f"Autores: {metadatos.get('authors', 'No disponible')}")
    doc.add_paragraph(f"Revista: {metadatos.get('journal', 'No disponible')}")
    doc.add_paragraph(f"Indexado en Scopus: {metadatos.get('is_scopus', 'No disponible')}")
    doc.add_paragraph(f"Cuartil SCImago: {metadatos.get('quartile', 'No disponible')}")
    doc.add_paragraph(f"País: {metadatos.get('country', 'No disponible')}")
    doc.add_paragraph(f"Área temática: {metadatos.get('subject_area', 'No disponible')}")
    doc.add_paragraph(f"Categoría temática: {metadatos.get('subject_category', 'No disponible')}")
    doc.add_paragraph(f"Entropía del texto generado: {entropia}")
    doc.add_paragraph()

    # Contenido del estado del arte
    doc.add_heading("Texto generado", level=2)
    for line in estado_arte_texto.split('\n'):
        doc.add_paragraph(line)

    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc.save(path)
    return path

# Obtiene información SCImago según el nombre de la revista
def get_scimago_info(journal_name):
    try:
        result = scimago_df[scimago_df['Title'].str.lower() == journal_name.lower()]
        if not result.empty:
            row = result.iloc[0]
            return {
                "quartile": row['Quartile'],
                "country": row['Country'],
                "subject_area": row['Areas'],
                "subject_category": row['Categories']
            }
        return {
            "quartile": "No encontrado",
            "country": "No disponible",
            "subject_area": "No disponible",
            "subject_category": "No disponible"
        }
    except Exception as e:
        print(f"Error buscando info SCImago: {e}")
        return {
            "quartile": "Error",
            "country": "Error",
            "subject_area": "Error",
            "subject_category": "Error"
        }

# Consulta los metadatos del artículo en Crossref si Elsevier falla
def get_crossref_metadata(doi):
    try:
        url = f"https://api.crossref.org/works/{doi}"
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            data = response.json()['message']
            title = data.get('title', ['Título no disponible'])[0]
            authors = ', '.join([f"{a.get('given', '')} {a.get('family', '')}" for a in data.get('author', [])])
            journal = data.get('container-title', ['Revista no disponible'])[0]
            scimago_info = get_scimago_info(journal)
            return {
                "title": title,
                "authors": authors,
                "journal": journal,
                "is_scopus": "No disponible",
                "quartile": scimago_info["quartile"],
                "country": scimago_info["country"],
                "subject_area": scimago_info["subject_area"],
                "subject_category": scimago_info["subject_category"]
            }
        else:
            print("DOI no encontrado en Crossref")
            return None
    except Exception as e:
        print(f"Error en Crossref: {e}")
        return None

# Consulta metadatos del artículo desde Elsevier y usa fallback Crossref
def get_article_details(doi):
    url = f"https://api.elsevier.com/content/article/doi/{doi}"
    headers = {'Accept': 'application/json', 'X-ELS-APIKey': API_KEY}
    response = requests.get(url, headers=headers)

    try:
        data = response.json()
        print("Respuesta Elsevier:", data)

        coredata = data.get('full-text-retrieval-response', {}).get('coredata', {})
        title = coredata.get('dc:title')

        if not title:
            print("Elsevier no tiene el recurso, usando Crossref...")
            return get_crossref_metadata(doi)

        authors_list = coredata.get('dc:creator', [])
        if isinstance(authors_list, list):
            authors = ', '.join(author.get('$', '') for author in authors_list)
        else:
            authors = authors_list.get('$', 'Autor no disponible')

        journal = coredata.get('prism:publicationName', 'Revista no disponible')
        is_scopus = "Sí" if 'scopus-id' in data.get('full-text-retrieval-response', {}) else "No"
        scimago_info = get_scimago_info(journal)

        return {
            "title": title,
            "authors": authors,
            "journal": journal,
            "is_scopus": is_scopus,
            "quartile": scimago_info["quartile"],
            "country": scimago_info["country"],
            "subject_area": scimago_info["subject_area"],
            "subject_category": scimago_info["subject_category"]
        }

    except Exception as e:
        print("Error procesando metadatos Elsevier:", e)
        return get_crossref_metadata(doi)

# Calcula la entropía de Shannon para medir diversidad léxica del texto
def calcular_entropia(texto):
    palabras = texto.split()
    total = len(palabras)
    if total == 0:
        return 0.0
    frecuencias = Counter(palabras)
    probabilidades = [f / total for f in frecuencias.values()]
    entropia = -sum(p * math.log2(p) for p in probabilidades)
    return round(entropia, 4)

# Ruta principal de carga y análisis del PDF
@app.route('/upload_pdf', methods=['POST'])
def upload_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No se ha enviado ningún archivo'}), 400

    file = request.files['file']
    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'error': 'Archivo inválido. Solo se permiten PDFs.'}), 400

    filename = secure_filename(file.filename)
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(path)

    texto = extract_text_from_pdf(path)
    if not texto:
        return jsonify({'error': 'No se pudo extraer texto del PDF.'}), 400

    estado = generate_estado_del_arte(texto)
    entropia = calcular_entropia(estado)
    doi = extract_doi_from_text(texto) or "10.1016/j.default"
    metadatos = get_article_details(doi)

    if metadatos is None:
        return jsonify({'error': 'No se pudo obtener metadatos del artículo.'}), 500

    nombre_word = f"estado_arte_{uuid.uuid4().hex[:8]}.docx"
    ruta_word = save_to_word(estado, nombre_word, metadatos, entropia)

    return jsonify({
        "title": metadatos["title"],
        "authors": metadatos["authors"],
        "journal": metadatos["journal"],
        "is_scopus": metadatos["is_scopus"],
        "quartile": metadatos["quartile"],
        "country": metadatos["country"],
        "subject_area": metadatos["subject_area"],
        "subject_category": metadatos["subject_category"],
        "estado_del_arte": estado,
        "entropia_estado_del_arte": entropia,
        "word_download_url": f"https://jofech20-github-io.onrender.com/download/{nombre_word}"
    }), 200

# Ruta de descarga de archivos Word
@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    return jsonify({'error': 'Archivo no encontrado.'}), 404

# Punto de entrada de la aplicación
if __name__ == '__main__':
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
